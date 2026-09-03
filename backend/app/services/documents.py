from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from uuid import UUID, uuid4

from docx import Document as DocxDocument
from fastapi import UploadFile
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings
from app.core.errors import AppError
from app.models.documents import (
    DocumentCategory,
    DocumentChunk,
    DocumentChunkEmbedding,
    DocumentSemanticStatus,
    DocumentStatus,
    ProjectDocument,
)
from app.models.project import Project
from app.schemas.documents import DocumentUpdate
from app.services.audit import AuditService
from app.services.authorization import AuthorizationService, Capability, accessible_project_ids
from app.services.projects import ProjectService
from app.storage import LocalDocumentStorage, create_document_storage

ALLOWED = {"pdf", "docx", "xlsx", "csv", "txt", "png", "jpg", "jpeg", "webp"}
EXTRACTABLE = {"pdf", "docx", "xlsx", "csv", "txt"}
STOP_WORDS = {
    "the",
    "and",
    "for",
    "with",
    "this",
    "that",
    "are",
    "from",
    "what",
    "come",
    "della",
    "delle",
    "degli",
    "che",
    "per",
    "con",
    "una",
    "uno",
    "del",
    "dei",
}


def _tokens(value: str) -> set[str]:
    return {word for word in re.findall(r"[\w-]{3,}", value.lower()) if word not in STOP_WORDS}


def _chunk_text(parts: list[tuple[str, dict]]) -> list[tuple[str, dict]]:
    chunks: list[tuple[str, dict]] = []
    for text, location in parts:
        clean = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        start = 0
        while start < len(clean) and len(chunks) < 300:
            end = min(len(clean), start + 1200)
            if end < len(clean):
                split = clean.rfind(" ", start + 700, end)
                if split > start:
                    end = split
            value = clean[start:end].strip()
            if value:
                chunks.append((value, location))
            if end >= len(clean):
                break
            start = max(start + 1, end - 150)
    return chunks


def extract_document(data: bytes, extension: str) -> list[tuple[str, dict]]:
    if extension == "txt":
        return [(data.decode("utf-8-sig", errors="replace")[:200_000], {"section": "text"})]
    if extension == "pdf":
        reader = PdfReader(io.BytesIO(data))
        return [
            ((page.extract_text() or "")[:30_000], {"page": index + 1})
            for index, page in enumerate(reader.pages[:50])
        ]
    if extension == "docx":
        document = DocxDocument(io.BytesIO(data))
        text = [paragraph.text for paragraph in document.paragraphs[:2000]]
        for table in document.tables[:50]:
            text.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows[:100])
        return [("\n".join(text)[:200_000], {"section": "document"})]
    if extension == "csv":
        reader = csv.reader(io.StringIO(data.decode("utf-8-sig", errors="replace")))
        rows = [" | ".join(row[:40]) for _, row in zip(range(101), reader, strict=False)]
        return [("\n".join(rows), {"sheet": "CSV"})]
    if extension == "xlsx":
        book = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        result = []
        for sheet in book.worksheets[:5]:
            rows = []
            for row in list(sheet.iter_rows(values_only=True))[:101]:
                rows.append(" | ".join("" if cell is None else str(cell) for cell in row[:30]))
            result.append(("\n".join(rows), {"sheet": sheet.title}))
        return result
    return []


class DocumentService:
    def __init__(self, session: AsyncSession, owner_user_id: UUID, settings: Settings) -> None:
        self.session = session
        self.owner_user_id = owner_user_id
        self.settings = settings
        self.audit = AuditService(session, owner_user_id)
        self.storage = create_document_storage(settings)

    async def _project(self, project_id: UUID, *, mutable: bool = False) -> Project:
        project = await ProjectService(self.session, self.owner_user_id).get(project_id)
        if mutable:
            ProjectService._ensure_mutable(project)
        return project

    async def _document(self, project_id: UUID, document_id: UUID) -> ProjectDocument:
        result = await self.session.execute(
            select(ProjectDocument)
            .join(Project)
            .where(
                ProjectDocument.id == document_id,
                ProjectDocument.project_id == project_id,
                Project.id.in_(accessible_project_ids(self.owner_user_id)),
            )
        )
        document = result.scalar_one_or_none()
        if document is None:
            raise AppError(
                code="document_not_found", message="Document not found.", status_code=404
            )
        if document.category == DocumentCategory.FINANCE:
            await AuthorizationService(self.session, self.owner_user_id).require(
                project_id, Capability.FINANCE_READ
            )
        return document

    async def list(self, project_id: UUID) -> list[ProjectDocument]:
        await self._project(project_id)
        query = select(ProjectDocument).where(ProjectDocument.project_id == project_id)
        if not await AuthorizationService(self.session, self.owner_user_id).can(
            project_id, Capability.FINANCE_READ
        ):
            query = query.where(ProjectDocument.category != DocumentCategory.FINANCE)
        result = await self.session.execute(query.order_by(ProjectDocument.created_at.desc()))
        return list(result.scalars())

    async def upload(
        self,
        project_id: UUID,
        upload: UploadFile,
        category: DocumentCategory,
        description: str | None,
    ) -> ProjectDocument:
        await self._project(project_id, mutable=True)
        if category == DocumentCategory.FINANCE:
            await AuthorizationService(self.session, self.owner_user_id).require(
                project_id, Capability.FINANCE_MANAGE
            )
        original = Path(upload.filename or "document").name[:255]
        extension = Path(original).suffix.lower().lstrip(".")
        if extension not in ALLOWED:
            raise AppError(
                code="document_type_not_supported",
                message="Unsupported document type.",
                status_code=415,
            )
        limit = self.settings.document_max_upload_mb * 1024 * 1024
        content = await upload.read(limit + 1)
        if len(content) > limit:
            raise AppError(
                code="document_too_large",
                message="Document exceeds the upload limit.",
                status_code=413,
            )
        if not content:
            raise AppError(
                code="document_empty", message="The uploaded document is empty.", status_code=422
            )
        document_id = uuid4()
        internal = f"{document_id.hex}.{extension}"
        relative = Path(str(project_id)) / internal
        storage_key = relative.as_posix()
        self.storage.put(storage_key, content, upload.content_type or "application/octet-stream")
        document = ProjectDocument(
            id=document_id,
            project_id=project_id,
            original_filename=original,
            internal_filename=internal,
            file_type=extension,
            mime_type=(upload.content_type or "application/octet-stream")[:120],
            size_bytes=len(content),
            category=category,
            description=description,
            storage_key=storage_key,
            status=DocumentStatus.PROCESSING
            if extension in EXTRACTABLE
            else DocumentStatus.UNSUPPORTED,
            processing_error=None
            if extension in EXTRACTABLE
            else "Text extraction is not available for this file type.",
            created_by_user_id=self.owner_user_id,
        )
        self.session.add(document)
        await self.session.flush()
        try:
            if extension in EXTRACTABLE:
                chunks = _chunk_text(extract_document(content, extension))
                for index, (text, location) in enumerate(chunks):
                    self.session.add(
                        DocumentChunk(
                            document_id=document.id,
                            project_id=project_id,
                            chunk_index=index,
                            text=text,
                            location=location,
                        )
                    )
                document.status = DocumentStatus.READY
                document.semantic_status = DocumentSemanticStatus.LEXICAL_ONLY
                document.processing_error = None
        except Exception:
            document.status = DocumentStatus.FAILED
            document.processing_error = (
                "Text extraction failed; the original file remains available."
            )
        self.audit.record(
            project_id=project_id,
            action="document.uploaded",
            entity_type="project_document",
            entity_id=document.id,
            changes={
                "filename": original,
                "file_type": extension,
                "size_bytes": len(content),
                "status": document.status.value,
            },
        )
        await self.session.commit()
        return await self._document(project_id, document.id)

    def path_for(self, document: ProjectDocument) -> Path:
        if not isinstance(self.storage, LocalDocumentStorage):
            raise RuntimeError("A local path is unavailable for object storage")
        return self.storage.path_for(document.storage_key)

    def open_file(self, document: ProjectDocument):
        return self.storage.open(document.storage_key)

    async def update(
        self, project_id: UUID, document_id: UUID, data: DocumentUpdate
    ) -> ProjectDocument:
        await self._project(project_id, mutable=True)
        document = await self._document(project_id, document_id)
        changes = data.model_dump(exclude_unset=True)
        if (
            document.category == DocumentCategory.FINANCE
            or changes.get("category") == DocumentCategory.FINANCE
        ):
            await AuthorizationService(self.session, self.owner_user_id).require(
                project_id, Capability.FINANCE_MANAGE
            )
        for key, value in changes.items():
            setattr(document, key, value)
        self.audit.record(
            project_id=project_id,
            action="document.updated",
            entity_type="project_document",
            entity_id=document.id,
            changes={"fields": list(changes)},
        )
        await self.session.commit()
        return await self._document(project_id, document_id)

    async def delete(self, project_id: UUID, document_id: UUID) -> None:
        await self._project(project_id, mutable=True)
        document = await self._document(project_id, document_id)
        self.audit.record(
            project_id=project_id,
            action="document.deleted",
            entity_type="project_document",
            entity_id=document.id,
            changes={"filename": document.original_filename},
        )
        await self.session.execute(
            delete(DocumentChunkEmbedding).where(
                DocumentChunkEmbedding.document_id == document.id
            )
        )
        await self.session.delete(document)
        await self.session.commit()
        self.storage.delete(document.storage_key)
